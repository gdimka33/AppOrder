import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
import openpyxl
import os
import re


class WordTableProcessor:
    def __init__(self, root):
        self.root = root
        self.root.title("Обработка таблиц Word")
        self.root.geometry("800x600")
        self.data = []
        self.current_department = ""
        self.processed_rows = 0
        self.total_rows = 0
        self.create_widgets()

    def create_widgets(self):
        # Основные виджеты
        self.select_btn = ttk.Button(
            self.root,
            text="Выбрать Word файл",
            command=self.select_word_file
        )
        self.select_btn.pack(pady=20)

        self.status_label = ttk.Label(
            self.root,
            text="Статус: ожидание выбора файла"
        )
        self.status_label.pack(pady=5)

        self.progress = ttk.Progressbar(
            self.root,
            orient="horizontal",
            length=400,
            mode="determinate"
        )
        self.progress.pack(pady=5)

        # Виджет для отображения информации о процессе
        self.process_info = ttk.Label(
            self.root,
            text="",
            font=('Arial', 8)
        )
        self.process_info.pack(pady=5)

        # Кнопка закрытия программы
        self.close_btn = ttk.Button(
            self.root,
            text="Закрыть программу",
            command=self.root.destroy,
            state=tk.DISABLED
        )
        self.close_btn.pack(pady=20)

    def update_process_info(self):
        """Обновить информацию о процессе обработки"""
        info_text = f"Обработано: {self.processed_rows} из {self.total_rows} строк"
        self.process_info.config(text=info_text)
        self.root.update_idletasks()

    def normalize_text(self, text):
        """Привести текст к строчному виду"""
        if not text:
            return ""
        return text.lower()

    def replace_rank_abbreviations(self, text):
        """Заменить сокращения в званиях и должностях"""
        if not text:
            return ""
        
        text = re.sub(r'\bст\s?\.', 'старший', text, flags=re.IGNORECASE)
        text = re.sub(r'\bмл\s?\.', 'младший', text, flags=re.IGNORECASE)
        return text

    def remove_vn_sl(self, text):
        """Убрать из текста 'вн. сл.' в разных формах"""
        return re.sub(r'вн\.?\с*сл\.?', '', text, flags=re.IGNORECASE).strip()

    def clean_department_name(self, text):
        """Удаляет текст в скобках из названия подразделения и обрабатывает переносы"""
        # Сначала удаляем текст в скобках
        text = re.sub(r'\([^)]*\)', '', text)
        # Заменяем все виды переносов на пробел и убираем лишние пробелы
        return ' '.join(text.split())

    def select_word_file(self):
        file_path = filedialog.askopenfilename(
            title="Выберите Word файл",
            filetypes=(("Word files", "*.docx"), ("All files", "*.*"))
        )
        if file_path:
            self.status_label.config(text=f"Обработка файла: {os.path.basename(file_path)}")
            self.progress['value'] = 0
            self.processed_rows = 0
            self.root.update_idletasks()
            self.process_word_file(file_path)

    def process_word_file(self, file_path):
        try:
            doc = Document(file_path)
            self.data = []
            self.current_department = ""
            self.total_rows = sum(len(table.rows) for table in doc.tables)
            self.processed_rows = 0
            self.update_process_info()

            for table in doc.tables:
                for row in table.rows:
                    if self.is_merged_row(row):
                        # Очищаем название подразделения от текста в скобках
                        department_text = self.clean_department_name(row.cells[0].text.strip())
                        self.current_department = self.normalize_text(department_text)
                    else:
                        self.process_employee_row(row)
                    
                    self.processed_rows += 1
                    self.progress['value'] = (self.processed_rows / self.total_rows) * 100
                    self.update_process_info()

            self.save_to_excel(file_path)
        except Exception as e:
            self.process_info.config(text=f"Ошибка: {str(e)}", foreground='red')
            self.progress['value'] = 0
            self.close_btn.config(state=tk.NORMAL)

    def is_merged_row(self, row):
        if len(row.cells) == 1:
            return True
        first_cell_text = row.cells[0].text.strip()
        for cell in row.cells[1:]:
            if cell.text.strip() != first_cell_text:
                return False
        return True

    def process_employee_row(self, row):
        if len(row.cells) >= 3:
            position_rank_text = row.cells[0].text.strip()
            # Заменяем перенос строки на пробел в ФИО
            full_name = ' '.join(row.cells[2].text.strip().split())

            if not full_name:
                return

            parts = [p.strip() for p in position_rank_text.split('\n') if p.strip()]
            if len(parts) == 0:
                position, rank = "", ""
            elif len(parts) == 1:
                position, rank = parts[0], ""
            else:
                position, rank = parts[0], '\н'.join(parts[1:])

            position = self.replace_rank_abbreviations(position)
            position = self.normalize_text(position)
            
            rank = self.remove_vn_sl(rank)
            rank = self.replace_rank_abbreviations(rank)
            rank = self.normalize_text(rank)

            self.data.append({
                'department': self.current_department,
                'position': position,
                'rank': rank,
                'full_name': self.normalize_text(full_name)
            })

    def save_to_excel(self, word_file_path):
        try:
            directory = os.path.dirname(word_file_path)
            excel_file_path = os.path.join(directory, "ОФИЦЕРЫ.xlsx")
            
            counter = 1
            while os.path.exists(excel_file_path):
                excel_file_path = os.path.join(directory, f"ОФИЦЕРЫ_{counter}.xlsx")
                counter += 1

            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "ОФИЦЕРЫ"

            # Измененный порядок столбцов
            headers = ["ФИО", "Звание", "Должность", "Подразделение"]
            for col_num, header in enumerate(headers, 1):
                ws.cell(row=1, column=col_num, value=header)

            for row_num, item in enumerate(self.data, 2):
                ws.cell(row=row_num, column=1, value=item['full_name'])
                ws.cell(row=row_num, column=2, value=item['rank'])
                ws.cell(row=row_num, column=3, value=item['position'])
                ws.cell(row=row_num, column=4, value=item['department'])

            for col in ws.columns:
                max_length = 0
                column = col[0].column_letter
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = (max_length + 2) * 1.2
                ws.column_dimensions[column].width = adjusted_width

            wb.save(excel_file_path)
            self.progress['value'] = 100
            self.process_info.config(
                text=f"Обработка завершена. Сохранено в: {excel_file_path}",
                foreground='green'
            )
            self.close_btn.config(state=tk.NORMAL)
        except Exception as e:
            self.process_info.config(text=f"Ошибка сохранения: {str(e)}", foreground='red')
            self.progress['value'] = 0
            self.close_btn.config(state=tk.NORMAL)


if __name__ == "__main__":
    root = tk.Tk()
    app = WordTableProcessor(root)
    root.mainloop()